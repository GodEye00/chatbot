from flask import current_app
import os
import re
import zipfile
import PyPDF2
from docx import Document
from io import BytesIO
import tempfile
import chardet
import io


# Function to import text from a file
def import_text_from_file(file_path):
    """Imports text from a TXT or PDF file."""
    if file_path.endswith(".txt"):
        try:
            with open(file_path, "r") as file:
                text = file.read()
                text = re.sub(r'\n+', ' ', text)
                text = text.strip()
            return text
        except FileNotFoundError:
            current_app.logger.exception(f"Error: File not found: {file_path}")
            return None
    elif file_path.endswith(".pdf"):
        try:
            with open(file_path, "rb") as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                full_text = ""
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    text = text.strip()
                    full_text += text
                full_text = re.sub(r'\n+', ' ', full_text)
                return full_text
        except (PyPDF2.errors.PdfReadError, FileNotFoundError) as e:
            current_app.logger.exception(f"Error reading PDF file: {file_path}")
            current_app.logger.exception(f"Exception: {e}")
            return None
    else:
        current_app.logger.error(f"Error: Unsupported file type: {file_path}")
        return None

# Function to import text from all TXT, PDF, and DOCX files in a ZIP file
def process_uploaded_file(file):
    file_contents = ''
    if file.filename.endswith('.txt'):
        # Read TXT file
        text = file.stream.read().decode('utf-8')
        text = re.sub(r'\n+', ' ', text)
        text = text.strip()
        file_contents = text
    elif file.filename.endswith('.pdf'):
        # Read PDF file
        try:
            pdf_reader = PyPDF2.PdfReader(file.stream)
            pdf_text = ""
            for page in pdf_reader.pages:
                extracted_page = page.extract_text() + " "
                text = extracted_page.strip()
                pdf_text += text
            pdf_text = re.sub(r'\n+', ' ', pdf_text)
            file_contents += pdf_text.strip()
        except Exception as e:
            current_app.logger.error(f"Error reading PDF file: {e}")
            return None
    elif file.filename.endswith('.docx'):
        # Read DOCX file
        try:
            doc = Document(file.stream)
            docx_text = " ".join(paragraph.text for paragraph in doc.paragraphs)
            docx_text = re.sub(r'\n+', ' ', docx_text)
            file_contents += docx_text.strip()
        except Exception as e:
            current_app.logger.error(f"Error reading DOCX file: {e}")
            return None
    elif file.filename.endswith('.zip'):
        # Process ZIP file
        try:
            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                file.save(temp_file)
                temp_file_path = temp_file.name

            with zipfile.ZipFile(temp_file_path, 'r') as zip_ref:
                for filename in zip_ref.namelist():
                    with zip_ref.open(filename) as f:
                        if filename.lower().endswith('.pdf'):
                            try:
                                pdf_memory_file = BytesIO(f.read())
                                pdf_reader = PyPDF2.PdfReader(pdf_memory_file)
                                pdf_text = ""
                                for page in pdf_reader.pages:
                                    pdf_text += page.extract_text() + " "
                                pdf_text = re.sub(r'\n+', ' ', pdf_text)
                                file_contents += pdf_text.strip()
                            except Exception as e:
                                current_app.logger.error(f"Error reading PDF file {filename} in ZIP: {e}")
                                continue  # Skip this file
                        else:
                            # For non-PDF files inside the ZIP
                            byte_content = f.read()
                            detected_encoding = chardet.detect(byte_content)['encoding']
                            if detected_encoding:
                                try:
                                    content = byte_content.decode(detected_encoding)
                                except UnicodeDecodeError:
                                    current_app.logger.exception(f"Could not decode {filename} with detected encoding {detected_encoding}")
                                    continue  # Skip this file
                            else:
                                current_app.logger.error(f"Could not detect encoding for {filename}")
                                continue  # Skip this file
                            content = re.sub(r'\n+', ' ', content)
                            file_contents += content

            os.remove(temp_file_path)
        except zipfile.BadZipFile:
            current_app.logger.exception("Invalid zip file")
            return None
        except Exception as e:
            current_app.logger.exception(f"Error processing ZIP file: {e}")
            return None
    else:
        current_app.logger.error("Unsupported file type")
        return None
    if len(file_contents) > 0:
        return file_contents
    else:
        raise Exception("File content is empty")



def process_file_content(file_content, filename=None):
    """
    Process file content based on the file extension and return the text.
    Supports direct bytes content for .txt, .pdf, .docx files, and a dictionary of these file types.
    :param file_content: Bytes content of the file or a dictionary containing multiple files.
    :param filename: The name of the file. Used if file_content is bytes. Ignored if file_content is a dictionary.
    """
    processed_text = ''

    if isinstance(file_content, dict):
        # If file_content is a dictionary, process each file within it
        for fname, content in file_content.items():
            processed_text += process_single_file(content, fname) + ' '
    else:
        # If file_content is not a dictionary, process it as a single file's bytes content
        processed_text = process_single_file(file_content, filename)

    return processed_text.strip()

def process_single_file(content, fname):
    """
    Process a single file's content based on its extension and return the text.
    """
    try:
        if fname.endswith('.txt'):
            detected_encoding = chardet.detect(content)['encoding']
            text = content.decode(detected_encoding)
            return re.sub(r'\n+', ' ', text).strip()

        elif fname.endswith('.pdf'):
            pdf_text = ''
            with io.BytesIO(content) as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                pdf_text = ' '.join(page.extract_text() for page in pdf_reader.pages if page.extract_text())
            return re.sub(r'\n+', ' ', pdf_text).strip()

        elif fname.endswith('.docx'):
            with io.BytesIO(content) as docx_file:
                doc = Document(docx_file)
                docx_text = ' '.join(paragraph.text for paragraph in doc.paragraphs)
            return re.sub(r'\n+', ' ', docx_text).strip()

        elif fname.endswith('.zip'):
            file_contents = ''
            with io.BytesIO(content) as zip_bio, zipfile.ZipFile(zip_bio) as zip_file:
                for zip_info in zip_file.infolist():
                    if not zip_info.is_dir() and zip_info.filename.endswith(('.txt', '.pdf', '.docx')):
                        with zip_file.open(zip_info) as extracted_file:
                            extracted_content = extracted_file.read()
                            content = process_single_file(extracted_content, zip_info.filename)
                            if content:
                                file_contents += content + ' '
            return file_contents.strip()

        else:
            current_app.logger.error(f"Unsupported file type for {fname}")
            return ''

    except Exception as e:
        current_app.logger.exception(f"Error processing file {fname}: {str(e)}")
        return ''

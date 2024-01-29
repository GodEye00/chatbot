from flask import current_app
import os
import json
import csv
import pandas as pd
from docx import Document
from reportlab.pdfgen import canvas

def write(file_data, ext, file_name, headers=None):
    """
    Writes data to a file with handling for different file extensions.

    Args:
    - file_data: The data to write. Can be a string, a list, or other types that can be converted to string.
    - ext: The extension or format of the file.
    - file_name: The name of the file to be created.
    - headers: The headers for CSV file (optional).

    Returns:
    - None
    """
    file_dir = "../docs/"
    file_path = os.path.join(file_dir, f"{file_name}.{ext}")

    # Ensure the directory exists
    os.makedirs(file_dir, exist_ok=True)

    current_app.logger.info(f"About to write to file: {file_path}")

    try:
        if ext == 'csv':
            if isinstance(file_data, list):
                with open(file_path, 'w', newline='', encoding='utf-8') as f:
                    writer = csv.writer(f)
                    if headers:
                        writer.writerow(headers)
                    writer.writerows(file_data)
            elif isinstance(file_data, pd.DataFrame):
                file_data.to_csv(file_path, index=False)
            else:
                raise ValueError("For CSV format, file_data must be a list of lists or a DataFrame.")
        elif ext == 'txt':
            with open(file_path, 'w', encoding='utf-8') as f:
                if isinstance(file_data, list):
                    file_data = '\n'.join(map(str, file_data))
                f.write(file_data)
        elif ext == 'json':
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(file_data, f, indent=4)
        elif ext == 'pdf':
            c = canvas.Canvas(file_path)
            if isinstance(file_data, list):
                for line in file_data:
                    c.drawString(72, 800, str(line))  # Basic PDF writing, adjust as needed
                    c.showPage()
            else:
                c.drawString(72, 800, str(file_data))
            c.save()
        elif ext == 'docx':
            doc = Document()
            if isinstance(file_data, list):
                for line in file_data:
                    doc.add_paragraph(str(line))
            else:
                doc.add_paragraph(str(file_data))
            doc.save(file_path)
        else:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(str(file_data))
    except IOError as e:
        current_app.logger.exception(f"Error writing to file: {e}")
